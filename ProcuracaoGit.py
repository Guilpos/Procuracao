import PySimpleGUI as sg
from docx import Document
from datetime import datetime
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class TelaPython:
    def __init__(self):
        # layout
        layout = [
            [sg.Text('Dados do Outorgante', size=(80, 0)), sg.Text('Dados do Procurador', size=(20, 0))],
            [sg.Text('Nome Completo', size=(14, 0)), sg.Input(key='nome'), sg.Text('Nome Completo', size=(15, 0)), sg.Input(key='nomeproc')],
            [sg.Text('Nacionalidade', size=(14, 0)), sg.Input(key='nacion'), sg.Text('Nacionalidade', size=(15, 0)), sg.Input(key='nacionproc')],
            [sg.Text('Estado Civil', size=(14, 0)), sg.Input(key='ec'), sg.Text('Estado Civil', size=(15, 0)), sg.Input(key='ecproc')],
            [sg.Text('Ocupação', size=(14, 0)), sg.Input(key='ocup'), sg.Text('Ocupação', size=(15, 0)), sg.Input(key='ocupproc')],
            [sg.Text('Nascimento', size=(14, 0)), sg.Input(key='nasc'), sg.Text('Nascimento.', size=(15, 0)), sg.Input(key='nascproc')],
            [sg.Text('RG', size=(14, 0)), sg.Input(key='rg'), sg.Text('RG ', size=(15, 0)), sg.Input(key='rgproc')],
            [sg.Text('Órgão Exp.', size=(14, 0)), sg.Input(key='oe'), sg.Text('Órgão Exp. Proc', size=(15, 0)), sg.Input(key='oeproc')],
            [sg.Text('CPF', size=(14, 0)), sg.Input(key='cpf'), sg.Text('CPF', size=(15, 0)), sg.Input(key='cpfproc')],
            [sg.Text('Cidade', size=(14, 0)), sg.Input(key='cidade'), sg.Text('Cidade', size=(15, 0)), sg.Input(key='cidadeproc')],
            [sg.Text('UF', size=(14, 0)), sg.Input(key='uf'), sg.Text('UF', size=(15, 0)), sg.Input(key='ufproc')],
            [sg.Text('Rua', size=(14, 0)), sg.Input(key='rua'), sg.Text('Rua', size=(15, 0)), sg.Input(key='ruaproc')],
            [sg.Text('Número', size=(14, 0)), sg.Input(key='numero'), sg.Text('Número', size=(15, 0)), sg.Input(key='numeroproc')],
            [sg.Text('Bairro', size=(14, 0)), sg.Input(key='bairro'), sg.Text('Bairro', size=(15, 0)), sg.Input(key='bairroproc')],
            [sg.Text('CEP', size=(14, 0)), sg.Input(key='cep'), sg.Text('CEP', size=(15, 0)), sg.Input(key='cepproc')],
            [sg.Text('Estado', size=(14, 0)), sg.Input('São Paulo', key='estado'), sg.Text('Estado', size=(15, 0)), sg.Input('São Paulo', key='estadoproc')],
            [sg.Text('Junta Com.', size=(14, 0)), sg.Input(key='jc')],
            [sg.Text('Mês', size=(14, 0)), sg.Input('Julho', key='mes')],
            [sg.Button('Enviar Dados')]

        ]

        # janela
        janela = sg.Window('Dados do Usuário').layout(layout)
        # Extrair os dados da tela
        self.button, self.values = janela.Read()

    def Iniciar(self):

        nome = self.values['nome']
        nacion = self.values['nacion']
        ec = self.values['ec']
        emp = self.values['ocup']
        nasc = self.values['nasc']
        rg = self.values['rg']
        oe = self.values['oe']
        cpf = self.values['cpf']
        cidade = self.values['cidade']
        uf = self.values['uf']
        rua = self.values['rua']
        num = self.values['numero']
        bairro = self.values['bairro']
        cep = self.values['cep']
        estado = self.values['estado']
        nomeproc = self.values['nomeproc']
        nacionproc = self.values['nacionproc']
        ecproc = self.values['ecproc']
        empproc = self.values['ocupproc']
        nascproc = self.values['nascproc']
        rgproc = self.values['rgproc']
        oeproc = self.values['oeproc']
        cpfproc = self.values['cpfproc']
        cidadeproc = self.values['cidadeproc']
        ufproc = self.values['ufproc']
        ruaproc = self.values['ruaproc']
        numproc = self.values['numeroproc']
        bairroproc = self.values['bairroproc']
        cepproc = self.values['cepproc']
        estadoproc = self.values['estadoproc']
        JC = self.values['jc']
        mes = self.values['mes']

        document = Document()
        paragraph1 = document.add_paragraph()
        paragraph1.add_run(f'{nome}').bold = True
        paragraph1.add_run(f', {nacion}, {ec}, ')
        if ec == 'casado' or ec == 'casada' or ec == 'Casado' or ec == 'Casada' or ec == 'CASADO' or ec == 'CASADA':
            paragraph1.add_run('Comunhão Parcial, ')
        paragraph1.add_run(f'{emp}, ')
        if nasc != '':
            paragraph1.add_run(f'nascido em {nasc}, ')
        paragraph1.add_run(f'portador da Cédula de Identidade RG nº {rg} - {oe}, n° do CPF:{cpf}, residente e domiciliado na cidade de {cidade} - {uf}, na {rua}, nº {num}, {bairro}, CEP: {cep}; ''')
        paragraph1.add_run('(“Outorgantes”) ').bold = True
        paragraph1.add_run('nomeia e constitui como seu bastante procurador: ')
        paragraph1.add_run(f'{nomeproc}').bold = True
        paragraph1.add_run(f', {nacionproc}, {ecproc}, ')
        if ecproc == 'casado' or ecproc == 'casada' or ecproc == 'Casado' or ecproc == 'Casada' or ecproc == 'CASADO' or ecproc == 'CASADA':
            paragraph1.add_run('Comunhão Parcial, ')
        paragraph1.add_run(f'{empproc}, nascido em {nascproc}, portador da Cédula de Identidade RG sob n°: {rgproc} {oeproc}, inscrito no CPF/MF sob n° {cpfproc}; e com o endereço comercial Cidade de {cidadeproc}, Estado de {estadoproc}, na {ruaproc}, {numproc}, {bairroproc} – {estadoproc}/{ufproc} – CEP: {cepproc} ')
        paragraph1.add_run('(“Outorgado”)').bold = True
        paragraph1.add_run(', para representar as outorgantes quando o recebimento de citações, notificações e intimações, em ações judiciais ou quaisquer medidas legais, seja de tribunais ordinários, especiais, arbitrais, e administrativos, que venham a ser instauradas contra as ')
        paragraph1.add_run('Outorgantes ').bold = True
        paragraph1.add_run('pela República Federativa do Brasil, relativas as ')
        paragraph1.add_run('Outorgantes').bold = True
        paragraph1.add_run(', em todo território nacional, perante (i) ')
        paragraph1.add_run(f'Junta Comercial do Estado de {JC}').underline = True
        paragraph1.add_run('; (ii) Secretarias da Fazenda; (iii) Receita Federal do Brasil; (iv) Instituto Nacional da Segurança Social “INSS”; (v) Prefeituras; e quais mais órgãos sejam necessários, com o fito especifico de inscrição/alteração/baixa de sociedades/empresas e sociedades, podendo para tanto assinar, tanto de forma física, quanto através de certificado digital, ')
        paragraph1.add_run('contratos e estatutos sociais, ').underline = True
        paragraph1.add_run('declarações de desimpedimentos – Art 1011 CC').bold = True
        paragraph1.add_run(', requerimentos, declarações, documentos fiscais, petições; pesquisar sobre situação fiscal ou cadastral, solicitar certidões, fichas cadastrais, documentos arquivados, emissão de guias, responsabilizando-se por todos os atos praticados no cumprimento deste instrumento, ao qual o outorgante faça parte do capital social, podendo, para tanto, substabelecer. O documento tem o prazo de 1 (um) ano a partir da data de assinatura.')
        paragraph1.add_run('')
        paragraph1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph1_format = paragraph1.paragraph_format

        paragraph2 = document.add_paragraph(' ')
        paragraph3 = document.add_paragraph(f'{estado}, {datetime.today().day} de {mes} de {datetime.today().year}')
        paragraph3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        paragraph3_format = paragraph3.paragraph_format
        paragraph3_format.line_spacing = Pt(10)
        paragraph3_format.space_after = Pt(80)

        paragraph4 = document.add_paragraph('________________________________________________________')
        paragraph4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph4_format = paragraph4.paragraph_format
        paragraph4_format.line_spacing = Pt(3)
        paragraph4_format.space_after = Pt(5)

        paragraph5 = document.add_paragraph(f'{nome}')
        paragraph5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph5_format = paragraph5.paragraph_format
        paragraph5_format.line_spacing = Pt(8)
        paragraph5_format.space_after = Pt(5)

        paragraph6 = document.add_paragraph(f'CPF Nº {cpf}')
        paragraph6.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        document.save(f'Procuração - {nome}.docx')


tela = TelaPython()
tela.Iniciar()
